import requests
from bs4 import BeautifulSoup
import csv
import os
from docx import Document

def scrape_website(url, csv_filename='scraped_data.csv', word_filename='scraped_data.docx'):
    try:
        # Send a GET request to the URL
        response = requests.get(url)
        response.raise_for_status()  # Raises an HTTPError for bad responses

        # Parse the HTML content using BeautifulSoup
        soup = BeautifulSoup(response.text, 'html.parser')

        # Extract all table rows (adjust selectors based on the structure)
        rows = soup.find_all('tr')

        # If rows are found, extract the data
        if rows:
            data = []
            # Loop through each row
            for row in rows:
                # Extract all cells within the row (th or td)
                cells = row.find_all(['th', 'td'])
                row_data = [cell.get_text(strip=True) for cell in cells]
                if row_data:  # Only add non-empty rows
                    data.append(row_data)

            # Get the path to the Documents folder
            documents_path = os.path.join(os.path.expanduser('~'), 'Documents')

            # Save the data to a CSV file in the Documents folder
            csv_file_path = os.path.join(documents_path, csv_filename)
            with open(csv_file_path, 'w', newline='', encoding='utf-8') as file:
                writer = csv.writer(file)
                writer.writerows(data)

            print(f"Data has been successfully scraped and saved to {csv_file_path}")

            # Convert CSV to Word document
            word_file_path = os.path.join(documents_path, word_filename)
            doc = Document()

            # Add a title to the Word document
            doc.add_heading('Scraped Data', 0)

            # Read CSV and write to Word document
            with open(csv_file_path, 'r', encoding='utf-8') as file:
                reader = csv.reader(file)
                for row in reader:
                    doc.add_paragraph(', '.join(row))  # Add each row as a paragraph

            # Save the Word document in the Documents folder
            doc.save(word_file_path)

            print(f"Data has been successfully converted and saved to {word_file_path}")
        else:
            print("No data found on the webpage.")

    except requests.exceptions.HTTPError as http_err:
        print(f"HTTP error occurred: {http_err}")
    except Exception as err:
        print(f"An error occurred: {err}")

# URL of the webpage to scrape
url = "https://nvd.nist.gov/vuln/vulnerability-detail-pages"  # Replace with the URL you want to scrape

# Call the function with the desired URL, output CSV filename, and output Word filename
scrape_website(url, 'Vulnerability_data08.csv', 'Vulnerability_data0df9.docx')
